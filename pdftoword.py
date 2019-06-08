from pdfminer.pdfparser import PDFParser
from pdfminer.pdfparser import PDFDocument
from pdfminer.pdfparser import PDFPage
from pdfminer.layout import LAParams,LTTextBoxHorizontal,LTImage,LTFigure
from pdfminer.pdfinterp import PDFResourceManager
from pdfminer.converter import PDFPageAggregator
from pdfminer.pdfinterp import PDFPageInterpreter, PDFTextExtractionNotAllowed
from docx import Document

path = input('pdf path:')
fp = open(path,'rb')
parse = PDFParser(fp)
# have password
# doc = PDFDocument(parse, password)
# no password
doc = PDFDocument()
parse.set_document(doc)
doc.set_parser(parse)
doc.initialize()
if not doc.is_extractable:
    raise PDFTextExtractionNotAllowed
else:
    rsrcmgr = PDFResourceManager()
    laparams = LAParams()
    device = PDFPageAggregator(rsrcmgr, laparams=laparams)
    interpreter = PDFPageInterpreter(rsrcmgr, device)
    document = Document()
    for page in doc.get_pages():
        interpreter.process_page(page)
        layout = device.get_result()
        print(layout)
        for x in layout:
            if isinstance(x, LTTextBoxHorizontal):
                #print(x.get_text().strip())
                document.add_paragraph(x.get_text().strip())
            elif isinstance(x, LTImage):
                pass
    document.save(r'C:\Users\lenovo\Desktop\test.docx')

